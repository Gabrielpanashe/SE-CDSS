"""Generate the Chapter 4 Execution Plan .docx document."""
import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

NAVY  = RGBColor(0x0F, 0x29, 0x44)
TEAL  = RGBColor(0x0E, 0xA5, 0xE9)
GREY  = RGBColor(0x64, 0x74, 0x8B)
BLACK = RGBColor(0x1E, 0x29, 0x3B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_bg(cell, hex_colour: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour)
    tcPr.append(shd)


def heading(text: str, size: int = 16, color=NAVY, bold: bool = True,
            sb: int = 12, sa: int = 6) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = color


def body(text: str, size: int = 11, color=BLACK, sb: int = 2,
         sa: int = 4, italic: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.italic = italic


def bullet(text: str, size: int = 11) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = BLACK


def divider() -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(chr(9472) * 75)
    r.font.size = Pt(8)
    r.font.color.rgb = GREY


# ── COVER PAGE ─────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
r = p.add_run("CHINHOYI UNIVERSITY OF TECHNOLOGY")
r.bold = True; r.font.size = Pt(13); r.font.color.rgb = NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Department of ICT and Electronics  |  BECE 5.2")
r.font.size = Pt(11); r.font.color.rgb = GREY

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("SE-CDSS")
r.bold = True; r.font.size = Pt(40); r.font.color.rgb = NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Sentiment-Enhanced Clinical Decision Support System for Precision Medicine")
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = TEAL

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("CHAPTER 4  |  EXECUTION PLAN")
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Evaluation, Results and System Upgrade Roadmap")
r.font.size = Pt(13); r.font.color.rgb = GREY; r.italic = True

doc.add_paragraph()
doc.add_paragraph()

tbl = doc.add_table(rows=4, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = "Table Grid"
for i, (label, value) in enumerate([
    ("Student",        "Panashe M. Chandiwana"),
    ("Student Number", "C21147799W"),
    ("Degree",         "B.Eng Honours  -  Computer Engineering"),
    ("Plan Date",      datetime.date.today().strftime("%d %B %Y")),
]):
    row = tbl.rows[i]
    set_cell_bg(row.cells[0], "0F2944")
    lr = row.cells[0].paragraphs[0].add_run(label)
    lr.bold = True; lr.font.color.rgb = WHITE; lr.font.size = Pt(10)
    vr = row.cells[1].paragraphs[0].add_run(value)
    vr.font.size = Pt(10); vr.font.color.rgb = BLACK

doc.add_page_break()

# ── 1. WHERE WE ARE ────────────────────────────────────────────────────────────
heading("1.  WHERE WE ARE RIGHT NOW", size=17, sb=6)
divider()
body("Chapters 1, 2, and 3 of the dissertation are fully written. The baseline "
     "system is complete and verified with 12 passing automated tests. The table "
     "below shows what exists now and what will be upgraded during Chapter 4.")
doc.add_paragraph()

tbl2 = doc.add_table(rows=6, cols=3)
tbl2.style = "Table Grid"
for i, hdr in enumerate(["Layer", "Current (Working Now)", "Chapter 4 Upgrade"]):
    set_cell_bg(tbl2.rows[0].cells[i], "0F2944")
    r = tbl2.rows[0].cells[i].paragraphs[0].add_run(hdr)
    r.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(10)

for i, (layer, cur, plan) in enumerate([
    ("Frontend",       "Streamlit  (dev/prototype only)",       "Next.js 14 + TypeScript + Tailwind CSS"),
    ("Backend API",    "FastAPI + Uvicorn",                     "FastAPI + Uvicorn  (no change needed)"),
    ("Database",       "SQLite via SQLAlchemy",                  "PostgreSQL via SQLAlchemy + Alembic"),
    ("Sentiment Model","TF-IDF + Logistic Regression",          "BioBERT  (dmis-lab/biobert-base-cased-v1.2)"),
    ("Deployment",     "Local script  (uvicorn ...)",           "Docker + docker-compose  (one command)"),
]):
    row = tbl2.rows[i + 1]
    for j, txt in enumerate([layer, cur, plan]):
        r = row.cells[j].paragraphs[0].add_run(txt)
        r.font.size = Pt(10); r.font.color.rgb = BLACK
        if j == 0: r.bold = True

doc.add_paragraph()
heading("What must NOT be touched without explicit instruction:", size=12, color=GREY,
        bold=True, sb=6, sa=2)
for item in [
    "src/  -  all preprocessing, model, recommendation, EHR, DB ORM code",
    "api/  -  FastAPI app and all 3 routes",
    "tests/  -  all 12 tests must remain green after every phase",
    "data/simulated_ehr/patients.json  -  50 synthetic patients, never modify",
    "src/config.py  -  all constants live here only",
]:
    bullet(item)

doc.add_page_break()

# ── 2. FOUR HYPOTHESES ────────────────────────────────────────────────────────
heading("2.  THE FOUR HYPOTHESES CHAPTER 4 MUST PROVE", size=17, sb=6)
divider()
body("Your dissertation states four research hypotheses. Every phase of work maps "
     "to answering at least one of them. Keep these visible as you work - they are "
     "the academic purpose behind each implementation decision.")
doc.add_paragraph()

for code, statement, evidence in [
    ("H1", "BioBERT achieves >= 80% macro F1-score on drug review sentiment classification.",
     "Answered by Phase D2  (BioBERT training + evaluation metrics)"),
    ("H2", "BioBERT outperforms TF-IDF + Logistic Regression on the same dataset and same split.",
     "Answered by Phase D1 + D2 together  (side-by-side comparison table)"),
    ("H3", "Combining sentiment + EHR history improves recommendation quality over structured data alone.",
     "Answered by Phase D3  (SHAP shows each factor's contribution to the recommendation)"),
    ("H4", "The system can be deployed in a low-resource environment with a single command.",
     "Answered by Phase G  (Docker + docker-compose)"),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    r1 = p.add_run(code + ":  ")
    r1.bold = True; r1.font.size = Pt(13); r1.font.color.rgb = TEAL
    r2 = p.add_run(statement)
    r2.font.size = Pt(11); r2.bold = True; r2.font.color.rgb = BLACK
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(1.2)
    r3 = p2.add_run("Evidence source: " + evidence)
    r3.font.size = Pt(10); r3.italic = True; r3.font.color.rgb = GREY

doc.add_page_break()

# ── 3. EXECUTION ORDER ────────────────────────────────────────────────────────
heading("3.  EXECUTION ORDER AND DEPENDENCY MAP", size=17, sb=6)
divider()
body("Phases run in the order below. Do not start a new phase until the current "
     "one is complete and all output files exist. D2 and D3 can overlap with E "
     "if you have available time.")
doc.add_paragraph()

phases_order = [
    ("D1", "Baseline Model Evaluation",   "2-3 hrs",   "H2 baseline",
     "Benchmark numbers that all comparisons are measured against."),
    ("D2", "BioBERT Sentiment Model",     "4-6 hrs",   "H1, H2",
     "Proves both hypotheses. Requires D1 for a fair split comparison."),
    ("D3", "SHAP Explainability",         "2 hrs",     "H3",
     "Fills src/explainability/. Adds word-level explanations to API."),
    ("E",  "PostgreSQL Migration",        "2-3 hrs",   "Architecture",
     "Replaces SQLite. Wires missing DB save calls. Can run parallel with D2/D3."),
    ("F",  "Next.js Frontend",            "8-12 hrs",  "Demonstration",
     "World-class UI replacing Streamlit. The dissertation demo artefact."),
    ("G",  "Docker Containerisation",     "2-3 hrs",   "H4",
     "Single-command deployment. Directly proves H4."),
]
tbl3 = doc.add_table(rows=len(phases_order) + 1, cols=5)
tbl3.style = "Table Grid"
for i, hdr in enumerate(["Phase", "Name", "Time", "Hypothesis", "Purpose"]):
    set_cell_bg(tbl3.rows[0].cells[i], "0F2944")
    r = tbl3.rows[0].cells[i].paragraphs[0].add_run(hdr)
    r.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(10)

for i, (code, name, time, hyp, purpose) in enumerate(phases_order):
    row = tbl3.rows[i + 1]
    set_cell_bg(row.cells[0], "0EA5E9")
    r0 = row.cells[0].paragraphs[0].add_run(code)
    r0.bold = True; r0.font.color.rgb = WHITE; r0.font.size = Pt(11)
    for j, txt in enumerate([name, time, hyp, purpose]):
        r = row.cells[j + 1].paragraphs[0].add_run(txt)
        r.font.size = Pt(10); r.font.color.rgb = BLACK

doc.add_page_break()

# ── 4. PHASE DETAILS ──────────────────────────────────────────────────────────
heading("4.  PHASE-BY-PHASE DETAILED PLAN", size=17, sb=6)
divider()

# D1
heading("PHASE D1  -  Baseline Model Evaluation", size=14, color=TEAL, sb=14)
body("GOAL: Produce dissertation Table 4.1 — the benchmark numbers all future "
     "comparisons are measured against.", italic=True)
body("What is this and why does it matter?", size=11, color=NAVY)
body(
    "The baseline TF-IDF + Logistic Regression model exists and works, but we "
    "have never formally measured it with proper statistics. In academic research, "
    "saying 'the model performs well' is meaningless without numbers. We run the "
    "model on data it has never seen before (the held-out test set) and measure "
    "objectively: how often is it right? How well does it handle each sentiment "
    "class separately? This produces the numbers that go in your results chapter."
)

heading("Steps (what we will build):", size=12, color=NAVY, sb=8, sa=2)
for s in [
    "Create scripts/evaluate_baseline.py",
    "Load the UCI Drug Review Dataset from data/raw/",
    "Perform an 80/20 stratified train/test split with random seed=42. "
    "Stratified means each sentiment class (positive, neutral, negative) is "
    "proportionally represented in both halves. Seed=42 means results are "
    "reproducible - anyone running the same code gets the same numbers.",
    "Apply the existing preprocessing pipeline (src/preprocessing/) to both halves",
    "Train TF-IDF + Logistic Regression on the training set only",
    "Run predictions on the test set only - the model has never seen this data",
    "Generate the classification report: accuracy, macro F1, per-class precision, recall, F1",
    "Save a confusion matrix heatmap PNG to docs/figures/baseline_confusion_matrix.png",
    "Save a per-class metrics bar chart PNG to docs/figures/baseline_per_class_metrics.png",
    "Save all metrics to docs/results/baseline_metrics.csv",
    "Print a clean summary table to the terminal",
]:
    bullet(s)

heading("Output files:", size=12, color=NAVY, sb=6, sa=2)
for f in [
    "scripts/evaluate_baseline.py",
    "docs/figures/baseline_confusion_matrix.png",
    "docs/figures/baseline_per_class_metrics.png",
    "docs/results/baseline_metrics.csv",
]:
    bullet(f)
body(
    "Definition of done: All output files exist. All 12 tests still pass. "
    "Numbers recorded in dissertation Table 4.1.",
    italic=True, color=GREY,
)

divider()

# D2
heading("PHASE D2  -  BioBERT Sentiment Model", size=14, color=TEAL, sb=14)
body("GOAL: Produce Table 4.2. Prove H1 (>= 80% macro F1) and H2 (BioBERT beats baseline).",
     italic=True)
body("What is BioBERT and why are we using it specifically?", size=11, color=NAVY)
body(
    "BioBERT is a version of Google's BERT model that was pre-trained on biomedical "
    "text - PubMed articles, clinical notes, drug literature. This means it already "
    "understands medical language, drug names, and clinical terminology before we "
    "even touch it. We then fine-tune it: we take those pre-trained weights and "
    "train a little more on our specific task (drug review sentiment classification). "
    "This is called transfer learning. It is much faster and cheaper than training "
    "from scratch, and it is standard practice in NLP research. For a clinical system "
    "like SE-CDSS, this domain knowledge is a genuine academic contribution - a "
    "general-purpose model does not understand the clinical meaning of words like "
    "'diuretic', 'contraindicated', or 'adverse event' the way BioBERT does."
)
body(
    "Hardware note: Your GPU has 7.9 GB VRAM (currently using only 1.2 GB). "
    "We use a 20,000-row stratified sample with batch_size=8 and max_length=128 "
    "tokens to stay within VRAM limits. Training takes approximately 2-3 hours.",
    color=GREY, italic=True,
)

heading("Steps:", size=12, color=NAVY, sb=8, sa=2)
for s in [
    "Create scripts/train_biobert.py",
    "Create a stratified 20,000-row sample from the UCI dataset using seed=42 "
    "(same seed as D1 for reproducibility)",
    "Apply the same 80/20 stratified split as D1 - this is critical for a fair comparison. "
    "If the splits differ, the comparison is scientifically invalid.",
    "Load dmis-lab/biobert-base-cased-v1.2 from HuggingFace Transformers library",
    "Fine-tune: batch_size=8, max_length=128 tokens, 3 epochs, learning_rate=2e-5",
    "Generate identical metrics to D1 (same format, same table structure, same test set)",
    "Save confusion matrix heatmap to docs/figures/biobert_confusion_matrix.png",
    "Save per-class metrics chart to docs/figures/biobert_per_class_metrics.png",
    "Save metrics to docs/results/biobert_metrics.csv",
    "Save the trained model to models/biobert_sentiment/ (already in .gitignore)",
    "Create src/models/predict_bert.py with the same interface as predict.py so "
    "the API can use either model without changes to the route code",
    "Create scripts/compare_models.py - loads both models, runs on the same test set, "
    "prints a side-by-side comparison table for your dissertation",
]:
    bullet(s)

heading("Output files:", size=12, color=NAVY, sb=6, sa=2)
for f in [
    "scripts/train_biobert.py",
    "src/models/predict_bert.py",
    "scripts/compare_models.py",
    "models/biobert_sentiment/",
    "docs/figures/biobert_confusion_matrix.png",
    "docs/figures/biobert_per_class_metrics.png",
    "docs/results/biobert_metrics.csv",
    "docs/results/model_comparison.csv",
]:
    bullet(f)

divider()

# D3
heading("PHASE D3  -  SHAP Explainability", size=14, color=TEAL, sb=14)
body("GOAL: Replace explanation: null in all API responses with real word-level explanations.",
     italic=True)
body("What is SHAP and why does it matter for a clinical system?", size=11, color=NAVY)
body(
    "SHAP (SHapley Additive exPlanations) is a technique from game theory that "
    "quantifies how much each input word contributed to the model's prediction. "
    "For example, if the model predicts 'Severe Adverse Reaction', SHAP can show "
    "that the word 'chest' pushed the score by -0.43 toward negative and the word "
    "'improved' pushed +0.12 toward positive. Your literature review already "
    "identified the black-box problem as a critical barrier to clinical adoption - "
    "a clinician will not trust a recommendation they cannot explain. SHAP is the "
    "solution. It makes the system transparent and clinically defensible. This "
    "directly addresses one of the research gaps you documented in Chapter 2."
)

heading("Steps:", size=12, color=NAVY, sb=8, sa=2)
for s in [
    "Create src/explainability/__init__.py  (the directory already exists but is empty)",
    "Create src/explainability/shap_explainer.py implementing get_explanation(text, top_n=5)",
    "get_explanation() returns a sorted list of {word, contribution} dicts "
    "ordered by absolute contribution value",
    "Wire the function into api/routes/feedback.py - replace explanation: null with real output",
    "Save a SHAP summary plot (all words, all classes) to docs/figures/shap_summary.png",
    "Save force plot examples for one positive and one negative review",
    "Add tests for the explainability module in tests/test_explainability.py",
]:
    bullet(s)

body(
    'The API explanation field changes from null to: '
    '[{"word": "pain", "contribution": -0.43}, {"word": "great", "contribution": +0.12}, ...]',
    size=10, color=GREY, italic=True,
)

divider()

# E
heading("PHASE E  -  PostgreSQL Migration", size=14, color=TEAL, sb=14)
body("GOAL: Replace SQLite with a production-grade local PostgreSQL database. "
     "Wire the currently missing DB save calls.", italic=True)
body("PostgreSQL Community Edition is completely free to install locally. No payment required.",
     color=GREY, italic=True)
body("What are we doing and why?", size=11, color=NAVY)
body(
    "SQLite is a single file on your disk - perfect for development but not suitable for "
    "production. It cannot handle multiple users writing at the same time (concurrent writes). "
    "PostgreSQL is the industry-standard free relational database. It runs locally on your "
    "machine, handles concurrent connections properly, and is consistent with what Chapter 3 "
    "documented as the planned migration. Additionally, right now the API routes compute "
    "recommendations and predictions but never save them to the database - the save_prediction() "
    "and save_recommendation() calls are missing from the routes. This phase connects those."
)

heading("Steps:", size=12, color=NAVY, sb=8, sa=2)
for s in [
    "Install PostgreSQL Community Edition  (free download from postgresql.org)",
    "Create local database se_cdss_db and user se_cdss_user",
    "Add psycopg2-binary and alembic to requirements.txt",
    "Create a .env file (never commit to git - already in .gitignore) "
    "containing: DATABASE_URL=postgresql://se_cdss_user:password@localhost/se_cdss_db",
    "Update src/config.py to read DATABASE_URL from the environment via python-dotenv",
    "Run: alembic init alembic  -  creates the migration infrastructure folder",
    "Configure alembic/env.py to point at the SQLAlchemy models in src/database/db.py",
    "Run: alembic revision --autogenerate  -  auto-generates SQL migration from the ORM models",
    "Run: alembic upgrade head  -  creates all 5 tables in PostgreSQL",
    "CRITICAL: Wire save_prediction() into api/routes/feedback.py  (currently missing)",
    "CRITICAL: Wire save_recommendation() into api/routes/recommend.py  (currently missing)",
    "Run all 12 tests to confirm nothing broke",
]:
    bullet(s)

divider()

# F
heading("PHASE F  -  Next.js Frontend", size=14, color=TEAL, sb=14)
body("GOAL: A world-class production UI replacing Streamlit. "
     "This is the dissertation demonstration artefact.", italic=True)
body("Why are we replacing Streamlit?", size=11, color=NAVY)
body(
    "Streamlit was always a development shortcut - fast to build but it looks like "
    "a developer tool, not a clinical product. Next.js 14 with TypeScript and "
    "Tailwind CSS is what professional healthcare software is built with. When you "
    "demonstrate this system to your examiners, the UI is the first thing they see. "
    "It needs to look credible, feel fast, and clearly communicate the clinical "
    "workflow to someone who may not be technical."
)

heading("Technology choices explained (so you understand each decision):", size=12,
        color=NAVY, sb=8, sa=2)
for name, explanation in [
    ("Next.js 14 App Router",
     "The leading React framework. App Router means each folder under /app is a "
     "URL route. Modern, fast, used in production healthcare apps worldwide."),
    ("TypeScript",
     "JavaScript with type declarations. The compiler catches bugs before you run "
     "the code. Standard in all professional frontend projects."),
    ("Tailwind CSS",
     "Write styles as class names directly in your HTML elements. Much faster than "
     "writing separate CSS files."),
    ("shadcn/ui",
     "Pre-built accessible components (buttons, cards, badges, forms). Professional "
     "look out of the box with no custom design work required."),
    ("Recharts",
     "React charting library. Used for the longitudinal sentiment trend chart on "
     "the clinician dashboard."),
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(3)
    r1 = p.add_run(name + ": ")
    r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = NAVY
    r2 = p.add_run(explanation)
    r2.font.size = Pt(10); r2.font.color.rgb = BLACK

heading("Project folder structure:", size=12, color=NAVY, sb=8, sa=2)
for f in [
    "frontend-next/app/page.tsx                  -  Landing / home page",
    "frontend-next/app/patient/page.tsx          -  Patient feedback portal",
    "frontend-next/app/clinician/page.tsx        -  Clinician dashboard",
    "frontend-next/components/FeedbackForm.tsx   -  Patient review submission form",
    "frontend-next/components/SentimentResult.tsx   -  Risk level + SHAP word chips",
    "frontend-next/components/RecommendationCard.tsx -  Drug card with score breakdown",
    "frontend-next/components/TrendChart.tsx     -  Recharts line chart (confidence over time)",
    "frontend-next/components/PatientLookup.tsx  -  Search patient by ID",
    "frontend-next/lib/api.ts                    -  All FastAPI calls centralised here",
    "frontend-next/types/index.ts                -  TypeScript interfaces for API responses",
]:
    bullet(f)

heading("Design requirements:", size=12, color=NAVY, sb=8, sa=2)
for r in [
    "Colour scheme: deep navy (#0F2944) + clinical white + accent teal (#0EA5E9)",
    "Risk badges: green (Mild Concern) / amber (Moderate Risk) / red (Severe Adverse Reaction)",
    "SHAP word contributions shown as coloured word chips - red = negative contribution, "
    "green = positive contribution",
    "Recommendation cards show the score breakdown: guideline score + EHR score + sentiment score",
    "Disclaimer displayed prominently on every result screen  (system rule - cannot be removed)",
    "Fully responsive - works on mobile and desktop",
    "Loading skeleton states on all async operations  (never blank screens)",
    "Error boundaries - user sees a friendly message if the API is down, not a crash page",
]:
    bullet(r)

divider()

# G
heading("PHASE G  -  Docker Containerisation", size=14, color=TEAL, sb=14)
body("GOAL: Prove H4. A single command starts the entire system - backend, frontend, and database.",
     italic=True)
body("What is Docker and why does this directly prove H4?", size=11, color=NAVY)
body(
    "Docker packages each part of your application into a self-contained container - "
    "a box that includes everything needed to run it (Python, Node.js, all dependencies, "
    "all configuration). docker-compose ties the boxes together and starts them in the "
    "correct order. This means anyone - a clinician, an examiner, a hospital IT department "
    "- can run your full system with one command without installing Python, Node.js, or "
    "PostgreSQL manually. This directly answers H4: if it runs with one command on any "
    "machine that has Docker, it is deployable in a low-resource environment."
)

heading("Files to create:", size=12, color=NAVY, sb=8, sa=2)
for f in [
    "backend/Dockerfile        -  packages the FastAPI backend into a container",
    "frontend-next/Dockerfile  -  packages the Next.js frontend into a container",
    "docker-compose.yml        -  ties backend + frontend + PostgreSQL together",
    ".env.example              -  safe template showing required variables (no real credentials)",
]:
    bullet(f)

heading("Requirements:", size=12, color=NAVY, sb=8, sa=2)
for r in [
    "docker compose up --build  starts the entire system from scratch in one command",
    "Backend accessible at localhost:8000  (FastAPI + OpenAPI docs at /docs)",
    "Frontend accessible at localhost:3000  (Next.js app)",
    "PostgreSQL data persisted via Docker volume  (data survives container restarts)",
    "Health checks on all services  (Docker waits for the database before starting the backend)",
    "README.md updated with a Docker deployment section",
]:
    bullet(r)

doc.add_page_break()

# ── 5. SYSTEM RULES ────────────────────────────────────────────────────────────
heading("5.  SYSTEM RULES  -  NEVER CHANGE THESE", size=17, sb=6)
divider()
body(
    "These rules are baked into the existing codebase and documented in Chapter 3. "
    "Breaking them creates a contradiction between your code and your dissertation. "
    "Every new file written during Chapter 4 must respect all of them."
)
doc.add_paragraph()

for name, rule, reason in [
    ("Recommendation formula",
     "Score = (0.33 x GuidelineScore) + (0.33 x EHRScore) + (0.33 x SentimentScore)",
     "Equal weighting reflects no clinical preference bias. Documented in Chapter 3."),
    ("Sentiment labels",
     "Always lowercase: positive, neutral, negative. Never capitalise.",
     "Consistency is enforced across the entire codebase and API contract."),
    ("Risk classification thresholds",
     "positive -> Mild Concern | neutral + conf>=0.60 -> Mild Concern | "
     "neutral + conf<0.60 -> Moderate Risk | negative + conf>=0.75 -> Severe Adverse Reaction | "
     "negative + conf<0.75 -> Moderate Risk",
     "These exact thresholds are in config.py. They are documented in Chapter 3 and must not change."),
    ("Keyword safety override",
     "[rash, chest pain, difficulty breathing, swelling, stopped taking, severe, emergency, "
     "hospital, reaction, allergic] -> always Severe Adverse Reaction regardless of model output.",
     "Safety rule applied before any ML prediction. This protects patients and cannot be overridden."),
    ("API response fields",
     "Every risk or recommendation response MUST include: disclaimer field + explanation field.",
     "Ethical and legal requirement for a clinical AI system. Both fields must always be present."),
    ("Five conditions only",
     "hypertension, diabetes, depression, malaria, respiratory",
     "System scope is intentionally limited. Do not add conditions without updating all drug maps."),
    ("Advisory only",
     "The system ASSISTS doctors - it never replaces them. "
     "Visible on every screen and in every API response.",
     "Core ethical principle of the entire research project."),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    r1 = p.add_run(name + ":  ")
    r1.bold = True; r1.font.color.rgb = NAVY; r1.font.size = Pt(11)
    r2 = p.add_run(rule)
    r2.font.size = Pt(11); r2.font.color.rgb = BLACK
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(1)
    r3 = p2.add_run("Why: " + reason)
    r3.font.size = Pt(10); r3.italic = True; r3.font.color.rgb = GREY

doc.add_page_break()

# ── 6. CODING STANDARDS ───────────────────────────────────────────────────────
heading("6.  CODING STANDARDS  -  EVERY NEW FILE", size=17, sb=6)
divider()
body(
    "Every new Python and TypeScript file created during Chapter 4 must follow these "
    "standards. They are not optional - they are what separates professional research "
    "software from a student script. Examiners will look at the code."
)
doc.add_paragraph()

for name, description in [
    ("Module docstring",
     "2-3 line comment at the very top of every Python file explaining what it does "
     "and why it exists."),
    ("Function docstrings",
     "Every function must have a docstring: what it does, what each parameter is, "
     "what it returns. No undocumented functions."),
    ("Type hints",
     "Every Python function signature declares input and output types. "
     "Example:  def predict(text: str) -> dict:"),
    ("No hardcoded paths",
     "Use os.path.join() from PROJECT_ROOT defined in config.py. "
     "Never write C:/Users/Panashe/... anywhere in code."),
    ("Constants in config.py only",
     "Thresholds, weights, labels, drug names - all defined in src/config.py. "
     "Never duplicated in any other file."),
    ("logging not print()",
     "Application code uses Python's logging module. "
     "print() is only acceptable in one-off evaluation scripts."),
    ("No raw tracebacks in API",
     "The API must never return a Python exception traceback to the user. "
     "Catch exceptions and return clean, readable error messages."),
    ("Confidence range",
     "Always a float between 0.0 and 1.0. Never a percentage. "
     "Never outside this range."),
    ("Disclaimer in every response",
     "The disclaimer field must be present in every risk or recommendation API response. "
     "This is a system rule, not a suggestion."),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    r1 = p.add_run(name + ": ")
    r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = TEAL
    r2 = p.add_run(description)
    r2.font.size = Pt(11); r2.font.color.rgb = BLACK

doc.add_page_break()

# ── 7. START INSTRUCTION ──────────────────────────────────────────────────────
heading("7.  START INSTRUCTION", size=17, sb=6)
divider()

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("Begin with Phase D1 only.")
r.bold = True; r.font.size = Pt(15); r.font.color.rgb = NAVY

body(
    "Do not touch any other phase until D1 is fully complete and all output files exist. "
    "Ask for confirmation before moving to D2. This is how we stay in control and make "
    "sure the dissertation numbers are correct before building on top of them."
)
doc.add_paragraph()
body("Phase D1 is complete when all of these files exist:")
for f in [
    "scripts/evaluate_baseline.py",
    "docs/figures/baseline_confusion_matrix.png",
    "docs/figures/baseline_per_class_metrics.png",
    "docs/results/baseline_metrics.csv",
]:
    bullet(f)
doc.add_paragraph()
body("And all 12 existing tests still pass:  pytest -v")
doc.add_paragraph()
body(
    "When D1 is confirmed, the results go directly into Table 4.1 of the dissertation. "
    "Then we confirm together and begin D2."
)
doc.add_paragraph()
doc.add_paragraph()
body(
    "This document is the single source of truth for all Chapter 4 development. "
    "Every decision, every file, and every phase follows this plan.",
    italic=True, color=GREY,
)

# ── SAVE ─────────────────────────────────────────────────────────────────────
out_path = "SE-CDSS Chapter 4 Execution Plan.docx"
doc.save(out_path)
print("Saved:", out_path)
